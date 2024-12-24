# Importaciones estándar de Python
import os
from concurrent.futures import ThreadPoolExecutor

# Importaciones de librerías de terceros
from flask import Flask, request, jsonify, render_template, send_file, url_for
import pandas as pd
from werkzeug.utils import secure_filename
from filelock import FileLock, Timeout

# Importaciones específicas para trabajar con documentos y PDFs
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
from pdfrw import PdfReader, PdfWriter, PageMerge
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

# Importaciones específicas para trabajar con Windows
import pythoncom
import win32com.client as win32


app = Flask(__name__)

# Carpetas para almacenar documentos
UPLOAD_FOLDER = "templates_word"  # Carpeta con las plantillas
OUTPUT_FOLDER = "output"     # Carpeta con los documentos generados
EXCEL_FILE = "output/data_global.xlsx"  # Archivo Excel global compartido
EXCEL_LOCK_FILE = "output/data_global.lock"  # Archivo de bloqueo para el Excel
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 # Limitar tamaño de archivo a 16 MB
executor = ThreadPoolExecutor(max_workers=2)

def replace_and_underline_text(paragraph, field, value):
    """
    Reemplaza un campo específico en un párrafo, sin preocuparse por el formato original.
    """
    if field in paragraph.text:
        # Reemplazar el texto directamente
        paragraph.text = paragraph.text.replace(field, value)

def replace_fields_in_table(table, data):
    """
    Reemplaza los campos dentro de una tabla en un documento Word.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for field, value in data.items():
                    replace_and_underline_text(paragraph, field, value)

def wait_for_excel_lock(lock_path, timeout=10):
    """
    Espera a que el archivo Excel esté desbloqueado.
    """
    lock = FileLock(lock_path)
    try:
        lock.acquire(timeout=timeout)
        return lock
    except Timeout:
        return jsonify({"warning": "El archivo Excel está siendo usado por otro proceso. Inténtalo nuevamente."}), 400


def save_to_excel(data):
    """
    Guarda los datos en un archivo Excel compartido.
    """
    try:
        # Esperar a que el archivo esté disponible
        with wait_for_excel_lock(EXCEL_LOCK_FILE):
            if os.path.exists(EXCEL_FILE):
                df = pd.read_excel(EXCEL_FILE)
                new_row = pd.DataFrame([data])
                df = pd.concat([df, new_row], ignore_index=True)
            else:
                df = pd.DataFrame([data])
            df.to_excel(EXCEL_FILE, index=False)
    except Exception as e:
        print(f"Error al guardar en Excel: {e}")


def save_pdf_from_docx(docx_path, pdf_path):
    try:
        # Inicializa COM antes de usar la biblioteca que interactúa con Word
        pythoncom.CoInitialize()  # Inicializa COM explícitamente
        
        # Usar docx2pdf para convertir el archivo
        convert(docx_path, pdf_path)
        
    except Exception as e:
        print(f"Error al convertir Word a PDF: {e}")
    finally:
        # Siempre hacer CoUninitialize() cuando termines
        pythoncom.CoUninitialize()  # Desinicializa COM

def process_word_generation(data, student_folder):
    try:
        docx_files = [f for f in os.listdir(UPLOAD_FOLDER) if f.endswith(".docx")]
        if not docx_files:
            return jsonify({"error": "No hay archivos Word en la carpeta seleccionada."}), 400

        for filename in docx_files:
            template_path = os.path.join(UPLOAD_FOLDER, filename)
            output_path = os.path.join(student_folder, f"output_{filename}")

            doc = Document(template_path)

            for paragraph in doc.paragraphs:
                for field, value in data.items():
                    replace_and_underline_text(paragraph, field, value)

            for table in doc.tables:
                replace_fields_in_table(table, data)

            doc.save(output_path)

            pdf_output_path = output_path.replace(".docx", ".pdf")
            save_pdf_from_docx(output_path, pdf_output_path)
        
        print(f"Formatos generados correctamente en la carpeta: {student_folder}")
        
        #return jsonify({"success": True, "message": "Formatos generados correctamente.", "output_folder": student_folder})
    except Exception as e:
        print(f"Error al generar documentos: {e}")
        #return jsonify({"error": "Error al generar documentos."}), 500

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_pdf_page_size(pdf_path):
    """Obtiene el tamaño de la primera página del PDF."""
    pdf = PdfReader(pdf_path)
    first_page = pdf.pages[0]
    media_box = first_page.MediaBox
    width = float(media_box[2]) - float(media_box[0])
    height = float(media_box[3]) - float(media_box[1])
    return width, height

def create_image_overlay(output_path, image_path, x, y, max_width, max_height, page_width, page_height):
    """Crea un PDF temporal con la imagen superpuesta."""
    c = canvas.Canvas(output_path, pagesize=(page_width, page_height))
    c.drawImage(ImageReader(image_path), x, y, max_width, max_height)
    c.save()

def merge_image_with_pdf(base_pdf_path, overlay_pdf_path, output_pdf_path):
    """Combina un PDF base con una imagen superpuesta preservando los formularios."""
    base_pdf = PdfReader(base_pdf_path)
    overlay_pdf = PdfReader(overlay_pdf_path)

    for base_page, overlay_page in zip(base_pdf.pages, overlay_pdf.pages):
        PageMerge(base_page).add(overlay_page).render()

    PdfWriter(output_pdf_path, trailer=base_pdf).write()


def generate_and_add_image_to_pdf(student_folder, image_path,pdf_image_path):
    """Genera documentos PDF y agrega una imagen a cada uno."""

    
    # Después de generar el PDF, añadir la imagen
    page_width, page_height = get_pdf_page_size(pdf_image_path)

    # Definir dimensiones máximas de la imagen
    max_width = 105
    max_height = (max_width / 3 * 4)

    # Crear el overlay de la imagen
    overlay_pdf_path = os.path.join(student_folder, "overlay_temp.pdf")
    create_image_overlay(overlay_pdf_path, image_path, x=457, y=755, max_width=max_width, max_height=max_height,
                            page_width=page_width, page_height=page_height)

    # Combinar el PDF original con la imagen
    merge_image_with_pdf(pdf_image_path, overlay_pdf_path, pdf_image_path)

    # Eliminar la imagen temporal después de agregarla
    os.remove(overlay_pdf_path)


@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate_words", methods=["POST"])
def generate_words():
    try:
        # Obtener datos fijos del formulario
        data = request.form.to_dict()


        # Obtener datos clave para el nombre de la carpeta
        grado = data.get("{{GRADO}}", "SIN_GRADO").replace(" ", "_")
        nombre = data.get("{{NOMBRE_EST}}", "SIN_NOMBRE").replace(" ", "_")
        apellido = data.get("{{APELLIDO_EST}}", "SIN_APELLIDO").replace(" ", "_")

        # Crear nombre de la carpeta en el formato requerido
        folder_name = f"{grado}_{nombre}_{apellido}"

        # Crear la carpeta específica para el estudiante
        student_folder = os.path.join(OUTPUT_FOLDER, folder_name)
        os.makedirs(student_folder, exist_ok=True)
        
        
        # Obtener archivo de imagen
        file = request.files.get('file')
        name, ext = os.path.splitext(file.filename)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(f"{folder_name}{ext}")
            img_path = os.path.join(student_folder, filename)
            file.save(img_path)
        else:
            return jsonify({"error": "Archivo de imagen no válido o no enviado."}), 400

        # Exportar datos a Excel
        save_to_excel(data)

        # Primero ejecutamos la generación de documentos en segundo plano
        future = executor.submit(process_word_generation, data, student_folder)

        # Esperamos a que la generación de documentos termine
        future.result()  # Esto bloqueará el hilo actual hasta que el proceso termine
        # Ahora que los documentos están generados, añadimos la imagen
        pdf_image_path = os.path.join(student_folder, "output_F03. MATRICULA.pdf")
        generate_and_add_image_to_pdf(student_folder, img_path,pdf_image_path)

        return jsonify({"success": True, "message": "Proceso iniciado. Espera mientras generamos los documentos.","output_folder": folder_name})
    except Exception as e:
        print(f"Error al procesar la solicitud: {e}")
        return jsonify({"error": "Error al procesar la solicitud."}), 500



@app.route("/preview/<folder_name>")
def preview_folder(folder_name):
    folder_path = os.path.join(OUTPUT_FOLDER, folder_name)
    if not os.path.exists(folder_path):
        return jsonify({"error": "La carpeta no existe."}), 404

    # Listar archivos en la carpeta
    files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
    file_links = [url_for('download_file', folder_name=folder_name, filename=f) for f in files]
    
    # Mostrar vista previa de los archivos generados
    return render_template("preview.html", folder_name=folder_name, files=zip(files, file_links))

@app.route("/download/<folder_name>/<filename>")
def download_file(folder_name, filename):
    file_path = os.path.join(OUTPUT_FOLDER, folder_name, filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "El archivo no existe."}), 404
    return send_file(file_path)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)


